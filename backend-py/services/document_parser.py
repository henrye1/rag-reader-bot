"""Document parsing with native Python + Gemini fallback for scanned PDFs.

Uses pdfplumber for PDF (with Gemini fallback for scanned docs),
python-docx for DOCX, direct decode for TXT/JSON.
"""

import json
import io
import os
import base64
import httpx


def parse_document(file_bytes: bytes, filename: str, mime_type: str) -> str:
    """Extract text from a document. Returns the extracted text string."""
    lower = filename.lower()

    # Plain text
    if mime_type == "text/plain" or lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")

    # JSON
    if mime_type == "application/json" or lower.endswith(".json"):
        content = file_bytes.decode("utf-8", errors="replace")
        try:
            parsed = json.loads(content)
            if isinstance(parsed, str):
                return parsed
            for key in ("text", "content", "prompt"):
                if key in parsed:
                    return str(parsed[key])
            return json.dumps(parsed, indent=2)
        except json.JSONDecodeError:
            raise ValueError("Invalid JSON file")

    # PDF
    if mime_type == "application/pdf" or lower.endswith(".pdf"):
        return _parse_pdf(file_bytes)

    # DOCX
    if lower.endswith(".docx") or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(file_bytes)

    raise ValueError(
        f"Unsupported file type: {mime_type}. "
        "Please use a plain text (.txt), JSON (.json), PDF (.pdf), or Word (.docx) file."
    )


def _parse_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

            # Also extract tables as text
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    cells = [str(c) if c else "" for c in row]
                    pages.append(" | ".join(cells))

    if pages:
        return "\n\n".join(pages)

    # Fallback: use Gemini to extract text from scanned/image-based PDFs
    print("[PDF] pdfplumber found no text — falling back to Gemini extraction")
    return _parse_pdf_with_gemini(file_bytes)


def _parse_pdf_with_gemini(file_bytes: bytes) -> str:
    """Use Gemini to extract text from scanned/image-based PDFs."""
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError("Could not extract text from PDF (scanned?) and GOOGLE_API_KEY not set for fallback.")

    b64 = base64.b64encode(file_bytes).decode("utf-8")
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"

    payload = {
        "contents": [{
            "parts": [
                {"inlineData": {"mimeType": "application/pdf", "data": b64}},
                {"text": "Extract ALL text content from this document. Return only the extracted text, preserving the original structure (headings, paragraphs, lists, tables). Do not add any commentary or explanation."},
            ]
        }],
        "generationConfig": {"temperature": 0.0},
    }

    resp = httpx.post(url, json=payload, timeout=120.0)
    if resp.status_code != 200:
        raise ValueError(f"Gemini PDF extraction failed ({resp.status_code}): {resp.text[:200]}")

    data = resp.json()
    try:
        text = data["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError):
        raise ValueError("Gemini returned no text for this PDF")

    if not text or len(text.strip()) < 10:
        raise ValueError("Could not extract meaningful text from PDF (scanned or empty)")

    print(f"[PDF] Gemini extracted {len(text)} characters")
    return text.strip()


def _parse_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    if not paragraphs:
        raise ValueError("Could not extract text from DOCX. The document may be empty.")

    return "\n\n".join(paragraphs)
